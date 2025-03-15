"""
Test script for file access components.
Use this to verify PDF, DOCX, and TXT file handling is working correctly.
"""
import os
import sys
import argparse
from pathlib import Path

# Add parent directory to path if running from subdirectory
parent_dir = os.path.dirname(os.path.abspath(__file__))
if parent_dir not in sys.path:
    sys.path.append(parent_dir)

def test_file_handlers():
    """Test the file handlers"""
    print("\n=== Testing File Handlers ===\n")

    # Import file handlers
    try:
        from file_handler import TextFileHandler, DocumentHandler, FileProcessor
        print("✓ Successfully imported file handlers")
    except ImportError as e:
        print(f"✗ Error importing file handlers: {str(e)}")
        return False

    # Initialize handlers
    file_dir = os.path.join("data", "files")
    doc_dir = os.path.join("data", "documents")

    # Create test directories if they don't exist
    os.makedirs(file_dir, exist_ok=True)
    os.makedirs(doc_dir, exist_ok=True)

    text_handler = TextFileHandler(file_dir)
    doc_handler = DocumentHandler(doc_dir)
    file_processor = FileProcessor(text_handler, doc_handler)

    print(f"✓ Initialized handlers for directories: {file_dir}, {doc_dir}")

    # Test PDF support
    print("\nTesting PDF support:")
    if hasattr(doc_handler, 'pdf_available') and doc_handler.pdf_available:
        print("✓ PyPDF library is available")

        # Create a test PDF file using a dummy method if no real PDF exists
        test_pdf_path = os.path.join(doc_dir, "test.pdf")
        if not os.path.exists(test_pdf_path):
            create_dummy_pdf(test_pdf_path)

        if os.path.exists(test_pdf_path):
            try:
                pdf_text = doc_handler.extract_text_from_pdf("test.pdf")
                print(f"✓ Successfully extracted text from PDF: {len(pdf_text)} characters")
                print(f"  Sample: {pdf_text[:100]}...")
            except Exception as e:
                print(f"✗ Error extracting text from PDF: {str(e)}")
        else:
            print("✗ No test PDF file available")
    else:
        print("✗ PyPDF library is not available - PDF support disabled")

    # Test DOCX support
    print("\nTesting DOCX support:")
    if hasattr(doc_handler, 'docx_available') and doc_handler.docx_available:
        print("✓ Python-docx library is available")

        # Create a test DOCX file using a dummy method if no real DOCX exists
        test_docx_path = os.path.join(doc_dir, "test.docx")
        if not os.path.exists(test_docx_path):
            create_dummy_docx(test_docx_path)

        if os.path.exists(test_docx_path):
            try:
                docx_text = doc_handler.extract_text_from_docx("test.docx")
                print(f"✓ Successfully extracted text from DOCX: {len(docx_text)} characters")
                print(f"  Sample: {docx_text[:100]}...")
            except Exception as e:
                print(f"✗ Error extracting text from DOCX: {str(e)}")
        else:
            print("✗ No test DOCX file available")
    else:
        print("✗ Python-docx library is not available - DOCX support disabled")

    # Test text file support
    print("\nTesting Text file support:")
    test_txt_path = os.path.join(file_dir, "test.txt")
    with open(test_txt_path, "w") as f:
        f.write("This is a test text file.\nIt has multiple lines.\nThis is used to test the text file handler.")

    try:
        txt_content = text_handler.read_text("test.txt")
        print(f"✓ Successfully read text file: {len(txt_content)} characters")
        print(f"  Sample: {txt_content[:50]}...")
    except Exception as e:
        print(f"✗ Error reading text file: {str(e)}")

    # Test file processor
    print("\nTesting file processor:")
    try:
        processed_files = file_processor.process_files_in_directory()
        file_types = {}
        for doc in processed_files:
            file_type = doc.get('type', 'unknown')
            if file_type not in file_types:
                file_types[file_type] = 0
            file_types[file_type] += 1

        print(f"✓ Successfully processed {len(processed_files)} files:")
        for file_type, count in file_types.items():
            print(f"  - {file_type}: {count} files")
    except Exception as e:
        print(f"✗ Error processing files: {str(e)}")

    return True

def create_dummy_pdf(output_path):
    """Create a dummy PDF file for testing"""
    try:
        from pypdf import PdfWriter, PdfReader
        from io import BytesIO

        # Check if we have a sample PDF to copy
        sample_path = os.path.join("resources", "sample.pdf")
        if os.path.exists(sample_path):
            import shutil
            shutil.copy(sample_path, output_path)
            return True

        # If no sample file, try to create a simple one
        writer = PdfWriter()
        page = writer.add_blank_page(width=612, height=792)

        # We can't easily add text directly with just pypdf,
        # but we can create a blank page
        with open(output_path, "wb") as f:
            writer.write(f)

        print(f"Created dummy PDF at {output_path} (blank page only)")
        return True
    except Exception as e:
        print(f"Could not create dummy PDF: {str(e)}")
        return False

def create_dummy_docx(output_path):
    """Create a dummy DOCX file for testing"""
    try:
        import docx

        # Check if we have a sample DOCX to copy
        sample_path = os.path.join("resources", "sample.docx")
        if os.path.exists(sample_path):
            import shutil
            shutil.copy(sample_path, output_path)
            return True

        # If no sample file, create a simple one
        doc = docx.Document()
        doc.add_heading("Test Document", 0)
        doc.add_paragraph("This is a test document created for testing the DOCX handling capabilities.")
        doc.add_paragraph("It contains multiple paragraphs to ensure proper text extraction.")

        # Add a table
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Header 1"
        table.cell(0, 1).text = "Header 2"
        table.cell(1, 0).text = "Value 1"
        table.cell(1, 1).text = "Value 2"

        doc.save(output_path)
        print(f"Created dummy DOCX at {output_path}")
        return True
    except Exception as e:
        print(f"Could not create dummy DOCX: {str(e)}")
        return False

def test_file_retrieval():
    """Test the file retrieval system"""
    print("\n=== Testing File Retrieval ===\n")

    # Import retrieval system
    try:
        from retrieval_system import KeywordRetriever, RetrievalContext
        print("✓ Successfully imported retrieval system")
    except ImportError as e:
        print(f"✗ Error importing retrieval system: {str(e)}")
        return False

    # Initialize retriever
    retriever = KeywordRetriever()
    context_manager = RetrievalContext(retriever)
    print("✓ Initialized retriever and context manager")

    # Import file processor
    try:
        from file_handler import TextFileHandler, DocumentHandler, FileProcessor

        file_dir = os.path.join("data", "files")
        doc_dir = os.path.join("data", "documents")

        text_handler = TextFileHandler(file_dir)
        doc_handler = DocumentHandler(doc_dir)
        file_processor = FileProcessor(text_handler, doc_handler)

        # Process and add documents to the retriever
        processed_files = file_processor.process_files_in_directory()
        for doc in processed_files:
            retriever.add_document(doc)

        print(f"✓ Added {len(processed_files)} documents to retriever")
    except Exception as e:
        print(f"✗ Error setting up document retrieval: {str(e)}")

    # Test retrieval
    try:
        # Test queries
        test_queries = [
            "test document",
            "text file",
            "multiple lines"
        ]

        for query in test_queries:
            context = context_manager.get_context_for_query(query)
            relevant_docs = context.get('relevant_documents', [])

            print(f"\nQuery: '{query}'")
            print(f"Found {len(relevant_docs)} relevant documents:")

            for i, doc in enumerate(relevant_docs):
                filename = doc.get('filename', 'Unknown')
                doc_type = doc.get('type', 'Unknown')
                score = doc.get('relevance_score', 0)

                print(f"  {i+1}. {filename} ({doc_type}) - Score: {score:.4f}")

                # Show a snippet of content
                content = doc.get('content', '')
                if isinstance(content, str):
                    snippet = content[:100] + "..." if len(content) > 100 else content
                    print(f"     Snippet: {snippet}")
                else:
                    print(f"     Content is not a string (type: {type(content)})")
    except Exception as e:
        print(f"✗ Error testing retrieval: {str(e)}")

    return True

def main():
    parser = argparse.ArgumentParser(description="Test file access components")
    parser.add_argument("--handlers", action="store_true", help="Test file handlers only")
    parser.add_argument("--retrieval", action="store_true", help="Test retrieval system only")
    args = parser.parse_args()

    print("Testing file access components...")

    if args.handlers or not (args.handlers or args.retrieval):
        test_file_handlers()

    if args.retrieval or not (args.handlers or args.retrieval):
        test_file_retrieval()

    print("\nTesting completed!")

    return 0

if __name__ == "__main__":
    sys.exit(main())