"""
File handler for AI Assistant.
Provides interfaces to access and process various file types.
"""
import os
import json
from typing import List, Dict, Any, Optional, Union, BinaryIO
import csv
from pathlib import Path
import re

class FileHandler:
    """Base class for file operations"""
    def __init__(self, base_directory: str = "data/files"):
        self.base_directory = base_directory
        # Create directory if it doesn't exist
        os.makedirs(base_directory, exist_ok=True)

    def get_full_path(self, filename: str) -> str:
        """Get full path for a file"""
        return os.path.join(self.base_directory, filename)

    def list_files(self, extension: str = None) -> List[str]:
        """List all files in the base directory, optionally filtered by extension"""
        files = []
        for file in os.listdir(self.base_directory):
            file_path = os.path.join(self.base_directory, file)
            if os.path.isfile(file_path):
                if extension is None or file.endswith(extension):
                    files.append(file)
        return sorted(files)

    def file_exists(self, filename: str) -> bool:
        """Check if a file exists"""
        return os.path.exists(self.get_full_path(filename))


class TextFileHandler(FileHandler):
    """Handler for text files (TXT, CSV, JSON)"""
    def read_text(self, filename: str) -> str:
        """Read text file content"""
        try:
            with open(self.get_full_path(filename), 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Fallback to binary mode if UTF-8 fails
            with open(self.get_full_path(filename), 'rb') as f:
                return f.read().decode('utf-8', errors='replace')

    def write_text(self, filename: str, content: str) -> bool:
        """Write content to text file"""
        try:
            with open(self.get_full_path(filename), 'w', encoding='utf-8') as f:
                f.write(content)
            return True
        except Exception as e:
            print(f"Error writing to {filename}: {str(e)}")
            return False

    def read_json(self, filename: str) -> Dict[str, Any]:
        """Read JSON file"""
        try:
            with open(self.get_full_path(filename), 'r', encoding='utf-8') as f:
                return json.load(f)
        except UnicodeDecodeError:
            # Fallback to binary mode if UTF-8 fails
            with open(self.get_full_path(filename), 'rb') as f:
                content = f.read().decode('utf-8', errors='replace')
                return json.loads(content)

    def write_json(self, filename: str, data: Dict[str, Any]) -> bool:
        """Write data to JSON file"""
        try:
            with open(self.get_full_path(filename), 'w', encoding='utf-8') as f:
                json.dump(data, f, indent=2)
            return True
        except Exception as e:
            print(f"Error writing JSON to {filename}: {str(e)}")
            return False

    def read_csv(self, filename: str, as_dict: bool = True) -> Union[List[Dict[str, Any]], List[List[str]]]:
        """Read CSV file"""
        try:
            with open(self.get_full_path(filename), 'r', encoding='utf-8', newline='') as f:
                if as_dict:
                    reader = csv.DictReader(f)
                    return list(reader)
                else:
                    reader = csv.reader(f)
                    return list(reader)
        except UnicodeDecodeError:
            # Fallback to binary mode with different encoding
            with open(self.get_full_path(filename), 'rb') as f:
                content = f.read()
                # Try common encodings
                for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                    try:
                        decoded = content.decode(encoding)
                        if as_dict:
                            reader = csv.DictReader(decoded.splitlines())
                            return list(reader)
                        else:
                            reader = csv.reader(decoded.splitlines())
                            return list(reader)
                    except:
                        continue
                # If all failed, use replacement character
                decoded = content.decode('utf-8', errors='replace')
                if as_dict:
                    reader = csv.DictReader(decoded.splitlines())
                    return list(reader)
                else:
                    reader = csv.reader(decoded.splitlines())
                    return list(reader)

    def write_csv(self, filename: str, data: List[Dict[str, Any]], fieldnames: List[str] = None) -> bool:
        """Write data to CSV file"""
        try:
            if not fieldnames and data:
                fieldnames = list(data[0].keys())

            with open(self.get_full_path(filename), 'w', encoding='utf-8', newline='') as f:
                writer = csv.DictWriter(f, fieldnames=fieldnames)
                writer.writeheader()
                writer.writerows(data)
            return True
        except Exception as e:
            print(f"Error writing CSV to {filename}: {str(e)}")
            return False


class DocumentHandler(FileHandler):
    """Handler for document files (PDF, DOCX)"""
    def __init__(self, base_directory: str = "data/documents"):
        super().__init__(base_directory)
        # Initialize document libraries if available
        self._init_pdf_library()
        self._init_docx_library()

    def _init_pdf_library(self):
        """Initialize PDF library if available"""
        try:
            import pypdf
            self.pdf_available = True
            print("PDF library loaded successfully")
        except ImportError:
            self.pdf_available = False
            print("PDF library not available - PDF support disabled")

    def _init_docx_library(self):
        """Initialize DOCX library if available"""
        try:
            import docx
            self.docx_available = True
            print("DOCX library loaded successfully")
        except ImportError:
            self.docx_available = False
            print("DOCX library not available - DOCX support disabled")

    def extract_text_from_pdf(self, filename: str) -> str:
        """Extract text from PDF file"""
        if not self.pdf_available:
            return f"PDF support not available. Cannot extract text from {filename}."

        try:
            from pypdf import PdfReader

            file_path = self.get_full_path(filename)
            reader = PdfReader(file_path)
            text = ""

            # Extract text from each page
            for page in reader.pages:
                text += page.extract_text() + "\n"

            # If no text was extracted, it might be a scanned PDF
            if not text.strip():
                return f"No extractable text found in {filename}. It may be a scanned document or image-based PDF."

            return text
        except Exception as e:
            print(f"Error extracting text from PDF {filename}: {str(e)}")
            return f"Error processing PDF: {str(e)}"

    def extract_text_from_docx(self, filename: str) -> str:
        """Extract text from DOCX file"""
        if not self.docx_available:
            return f"DOCX support not available. Cannot extract text from {filename}."

        try:
            import docx

            file_path = self.get_full_path(filename)
            doc = docx.Document(file_path)

            # Extract text from paragraphs
            paragraphs = [paragraph.text for paragraph in doc.paragraphs]

            # Extract text from tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    tables_text.append(" | ".join(row_text))

            # Combine all text
            all_text = "\n".join(paragraphs + ["\n"] + tables_text)
            return all_text
        except Exception as e:
            print(f"Error extracting text from DOCX {filename}: {str(e)}")
            return f"Error processing DOCX: {str(e)}"

    def get_document_metadata(self, filename: str) -> Dict[str, Any]:
        """Get metadata from document"""
        file_path = self.get_full_path(filename)
        stats = os.stat(file_path)

        metadata = {
            "filename": filename,
            "path": file_path,
            "size_bytes": stats.st_size,
            "created": stats.st_ctime,
            "modified": stats.st_mtime,
            "extension": Path(filename).suffix
        }

        # Extract additional metadata based on file type
        extension = Path(filename).suffix.lower()

        if extension == '.pdf' and self.pdf_available:
            try:
                from pypdf import PdfReader
                reader = PdfReader(file_path)

                # Add PDF-specific metadata
                metadata.update({
                    "page_count": len(reader.pages),
                    "pdf_info": reader.metadata
                })
            except:
                pass

        elif extension == '.docx' and self.docx_available:
            try:
                import docx
                doc = docx.Document(file_path)

                # Add DOCX-specific metadata
                metadata.update({
                    "paragraph_count": len(doc.paragraphs),
                    "table_count": len(doc.tables)
                })
            except:
                pass

        return metadata

    def is_supported_document(self, filename: str) -> bool:
        """Check if a file is a supported document type"""
        extension = Path(filename).suffix.lower()

        if extension == '.pdf':
            return self.pdf_available
        elif extension == '.docx':
            return self.docx_available

        return False


class FileProcessor:
    """Utility class for processing files for AI consumption"""
    def __init__(self, text_handler: TextFileHandler = None, doc_handler: DocumentHandler = None):
        self.text_handler = text_handler or TextFileHandler()
        self.doc_handler = doc_handler or DocumentHandler()

    def process_files_in_directory(self, directory: str = None) -> List[Dict[str, Any]]:
        """Process all supported files in a directory"""
        if directory:
            self.text_handler.base_directory = directory
            self.doc_handler.base_directory = directory

        results = []

        # Process text files
        for filename in self.text_handler.list_files(".txt"):
            try:
                content = self.text_handler.read_text(filename)
                results.append({
                    "filename": filename,
                    "content": content,
                    "type": "text"
                })
                print(f"Processed text file: {filename}")
            except Exception as e:
                print(f"Error processing text file {filename}: {str(e)}")

        # Process JSON files
        for filename in self.text_handler.list_files(".json"):
            try:
                content = self.text_handler.read_json(filename)
                results.append({
                    "filename": filename,
                    "content": content,
                    "type": "json"
                })
                print(f"Processed JSON file: {filename}")
            except Exception as e:
                print(f"Error processing JSON file {filename}: {str(e)}")

        # Process CSV files
        for filename in self.text_handler.list_files(".csv"):
            try:
                content = self.text_handler.read_csv(filename)
                results.append({
                    "filename": filename,
                    "content": content,
                    "type": "csv"
                })
                print(f"Processed CSV file: {filename}")
            except Exception as e:
                print(f"Error processing CSV file {filename}: {str(e)}")

        # Process PDF files
        if self.doc_handler.pdf_available:
            for filename in self.doc_handler.list_files(".pdf"):
                try:
                    content = self.doc_handler.extract_text_from_pdf(filename)
                    results.append({
                        "filename": filename,
                        "content": content,
                        "type": "pdf"
                    })
                    print(f"Processed PDF file: {filename}")
                except Exception as e:
                    print(f"Error processing PDF file {filename}: {str(e)}")

        # Process DOCX files
        if self.doc_handler.docx_available:
            for filename in self.doc_handler.list_files(".docx"):
                try:
                    content = self.doc_handler.extract_text_from_docx(filename)
                    results.append({
                        "filename": filename,
                        "content": content,
                        "type": "docx"
                    })
                    print(f"Processed DOCX file: {filename}")
                except Exception as e:
                    print(f"Error processing DOCX file {filename}: {str(e)}")

        return results

    def chunk_document(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
        """Split document into chunks for processing"""
        chunks = []
        start = 0

        while start < len(text):
            end = min(start + chunk_size, len(text))
            # Try to find a good break point (period, newline, etc.)
            if end < len(text):
                # Look for a period or newline within the last 100 chars of the chunk
                for i in range(end, max(end - 100, start), -1):
                    if text[i] in ['.', '\n', '!', '?']:
                        end = i + 1
                        break

            chunks.append(text[start:end])
            start = end - overlap  # Create overlap

        return chunks

    def get_file_type(self, filename: str) -> str:
        """Detect file type from filename"""
        extension = Path(filename).suffix.lower()

        if extension == '.txt':
            return 'text'
        elif extension == '.json':
            return 'json'
        elif extension == '.csv':
            return 'csv'
        elif extension == '.pdf':
            return 'pdf'
        elif extension == '.docx':
            return 'docx'
        else:
            return 'unknown'

    def process_file(self, filename: str) -> Dict[str, Any]:
        """Process a single file based on its type"""
        file_type = self.get_file_type(filename)

        if file_type == 'text':
            content = self.text_handler.read_text(filename)
            return {"filename": filename, "content": content, "type": "text"}

        elif file_type == 'json':
            content = self.text_handler.read_json(filename)
            return {"filename": filename, "content": content, "type": "json"}

        elif file_type == 'csv':
            content = self.text_handler.read_csv(filename)
            return {"filename": filename, "content": content, "type": "csv"}

        elif file_type == 'pdf' and self.doc_handler.pdf_available:
            content = self.doc_handler.extract_text_from_pdf(filename)
            return {"filename": filename, "content": content, "type": "pdf"}

        elif file_type == 'docx' and self.doc_handler.docx_available:
            content = self.doc_handler.extract_text_from_docx(filename)
            return {"filename": filename, "content": content, "type": "docx"}

        else:
            return {"filename": filename, "content": f"Unsupported file type: {file_type}", "type": "unknown"}