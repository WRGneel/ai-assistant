"""
Creates sample data files for testing the AI assistant.
This script generates text, PDF, DOCX, JSON, and CSV files.
"""
import os
import sys
import json
import csv
import random
import argparse
from datetime import datetime, timedelta

# Check if required libraries are available
try:
    from pypdf import PdfWriter
    pdf_available = True
except ImportError:
    pdf_available = False
    print("Warning: pypdf library not available. PDF generation disabled.")

try:
    import docx
    docx_available = True
except ImportError:
    docx_available = False
    print("Warning: python-docx library not available. DOCX generation disabled.")

def create_sample_text_files(output_dir, count=3):
    """Create sample text files"""
    print(f"Creating {count} sample text files in {output_dir}...")
    os.makedirs(output_dir, exist_ok=True)

    topics = [
        "artificial intelligence",
        "machine learning",
        "natural language processing",
        "computer vision",
        "neural networks",
        "deep learning",
        "robotics",
        "data science",
        "big data",
        "cloud computing"
    ]

    for i in range(count):
        topic = random.choice(topics)
        filename = f"{topic.replace(' ', '_')}_info.txt"
        file_path = os.path.join(output_dir, filename)

        with open(file_path, "w") as f:
            f.write(f"# Information about {topic.title()}\n\n")
            f.write(f"{topic.title()} is a field of computer science that focuses on creating intelligent machines.\n\n")
            f.write(f"This document provides an overview of {topic} concepts and applications.\n\n")
            f.write("## Key Concepts\n\n")
            f.write(f"1. {topic.title()} involves the development of algorithms and models.\n")
            f.write(f"2. Modern {topic} systems rely on large amounts of data.\n")
            f.write(f"3. {topic.title()} applications are used in various industries.\n\n")
            f.write("## Applications\n\n")
            f.write(f"- Healthcare: {topic} is used for diagnosis and treatment planning.\n")
            f.write(f"- Finance: {topic} helps with fraud detection and risk assessment.\n")
            f.write(f"- Manufacturing: {topic} enables predictive maintenance and quality control.\n")

        print(f"  Created {filename}")

    return True

def create_sample_pdf_files(output_dir, count=2):
    """Create sample PDF files"""
    if not pdf_available:
        print("Skipping PDF creation - pypdf library not available")
        return False

    print(f"Creating {count} sample PDF files in {output_dir}...")
    os.makedirs(output_dir, exist_ok=True)

    topics = [
        "quarterly report",
        "annual review",
        "market analysis",
        "research findings",
        "project proposal",
        "technical specification"
    ]

    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        reportlab_available = True
    except ImportError:
        reportlab_available = False
        print("Warning: reportlab library not available. Using basic PDF generation.")

    for i in range(count):
        topic = random.choice(topics)
        filename = f"{topic.replace(' ', '_')}.pdf"
        file_path = os.path.join(output_dir, filename)

        if reportlab_available:
            # Create PDF with text content
            c = canvas.Canvas(file_path, pagesize=letter)
            width, height = letter

            # Add title
            c.setFont("Helvetica-Bold", 16)
            title = f"{topic.title()} - {datetime.now().strftime('%B %Y')}"
            c.drawString(72, height - 72, title)

            # Add content
            c.setFont("Helvetica", 12)
            y = height - 100

            c.drawString(72, y, f"This document contains information about {topic}.")
            y -= 20

            c.drawString(72, y, "Executive Summary:")
            y -= 15

            c.setFont("Helvetica", 10)
            summary = (
                f"This {topic} provides an overview of performance metrics, "
                f"key achievements, and strategic objectives for the organization. "
                f"It highlights important data points and trends observed during "
                f"the reporting period and offers recommendations for future actions."
            )

            # Simple word wrapping
            words = summary.split()
            line = ""
            for word in words:
                test_line = line + word + " "
                line_width = c.stringWidth(test_line, "Helvetica", 10)

                if line_width < width - 144:  # 72pt margins on each side
                    line = test_line
                else:
                    c.drawString(72, y, line)
                    y -= 15
                    line = word + " "

            if line:
                c.drawString(72, y, line)

            c.save()
        else:
            # Create a simple blank PDF if reportlab is not available
            writer = PdfWriter()
            writer.add_blank_page(width=612, height=792)  # letter size

            with open(file_path, "wb") as f:
                writer.write(f)

        print(f"  Created {filename}")

    return True

def create_sample_docx_files(output_dir, count=2):
    """Create sample DOCX files"""
    if not docx_available:
        print("Skipping DOCX creation - python-docx library not available")
        return False

    print(f"Creating {count} sample DOCX files in {output_dir}...")
    os.makedirs(output_dir, exist_ok=True)

    topics = [
        "meeting minutes",
        "project plan",
        "technical documentation",
        "user guide",
        "policy document",
        "requirements specification"
    ]

    for i in range(count):
        topic = random.choice(topics)
        filename = f"{topic.replace(' ', '_')}.docx"
        file_path = os.path.join(output_dir, filename)

        doc = docx.Document()

        # Add title
        doc.add_heading(f"{topic.title()} - {datetime.now().strftime('%B %d, %Y')}", 0)

        # Add introduction
        doc.add_paragraph(f"This document contains {topic} information for reference purposes.")

        # Add content
        doc.add_heading("Overview", 1)
        doc.add_paragraph(
            f"This {topic} provides details about the project structure, "
            f"key components, and implementation guidelines. It serves as a "
            f"reference for team members and stakeholders."
        )

        # Add a list
        doc.add_heading("Key Points", 1)
        points = [
            "Regular team meetings are scheduled on Mondays",
            "Progress reports are due at the end of each sprint",
            "All code must pass automated tests before merging",
            "Documentation should be updated with each feature release"
        ]
        for point in points:
            p = doc.add_paragraph()
            p.style = 'List Bullet'
            p.add_run(point)

        # Add a table
        doc.add_heading("Schedule", 1)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'

        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Date"
        header_cells[1].text = "Milestone"
        header_cells[2].text = "Owner"

        # Add data rows
        start_date = datetime.now()
        milestones = [
            "Project Kickoff",
            "Requirements Analysis",
            "Design Phase",
            "Implementation",
            "Testing"
        ]
        owners = ["John", "Sarah", "Michael", "Emily", "David"]

        for j in range(5):
            milestone_date = start_date + timedelta(days=j*14)
            row_cells = table.add_row().cells
            row_cells[0].text = milestone_date.strftime("%Y-%m-%d")
            row_cells[1].text = milestones[j]
            row_cells[2].text = owners[j]

        doc.save(file_path)
        print(f"  Created {filename}")

    return True

def create_sample_json_files(output_dir, count=2):
    """Create sample JSON files"""
    print(f"Creating {count} sample JSON files in {output_dir}...")
    os.makedirs(output_dir, exist_ok=True)

    # Sample product catalog
    product_filename = "product_catalog.json"
    product_file_path = os.path.join(output_dir, product_filename)

    products = {
        "products": [
            {"id": 1, "name": "Laptop Pro", "price": 1299.99, "category": "Electronics", "in_stock": True},
            {"id": 2, "name": "Wireless Headphones", "price": 149.99, "category": "Electronics", "in_stock": True},
            {"id": 3, "name": "Smart Watch", "price": 249.99, "category": "Electronics", "in_stock": False},
            {"id": 4, "name": "Programming Guide", "price": 39.99, "category": "Books", "in_stock": True},
            {"id": 5, "name": "Coffee Maker", "price": 89.99, "category": "Kitchen", "in_stock": True},
            {"id": 6, "name": "Bluetooth Speaker", "price": 79.99, "category": "Electronics", "in_stock": True}
        ],
        "metadata": {
            "last_updated": datetime.now().strftime("%Y-%m-%d"),
            "currency": "USD",
            "total_products": 6
        }
    }

    with open(product_file_path, "w") as f:
        json.dump(products, f, indent=2)

    print(f"  Created {product_filename}")

    # Sample user data
    user_filename = "user_profiles.json"
    user_file_path = os.path.join(output_dir, user_filename)

    users = {
        "users": [
            {
                "id": 101,
                "name": "Alice Johnson",
                "email": "alice.j@example.com",
                "role": "Administrator",
                "active": True,
                "last_login": "2023-05-15T10:30:45Z"
            },
            {
                "id": 102,
                "name": "Bob Smith",
                "email": "bob.smith@example.com",
                "role": "User",
                "active": True,
                "last_login": "2023-05-14T08:12:22Z"
            },
            {
                "id": 103,
                "name": "Carol Davis",
                "email": "carol.d@example.com",
                "role": "Editor",
                "active": False,
                "last_login": "2023-04-28T14:45:33Z"
            }
        ],
        "metadata": {
            "total_users": 3,
            "active_users": 2,
            "database_version": "2.4.1"
        }
    }

    with open(user_file_path, "w") as f:
        json.dump(users, f, indent=2)

    print(f"  Created {user_filename}")

    return True

def create_sample_csv_files(output_dir, count=2):
    """Create sample CSV files"""
    print(f"Creating {count} sample CSV files in {output_dir}...")
    os.makedirs(output_dir, exist_ok=True)

    # Sales data
    sales_filename = "monthly_sales.csv"
    sales_file_path = os.path.join(output_dir, sales_filename)

    sales_data = [
        ["date", "product_id", "product_name", "quantity", "unit_price", "total"],
        ["2023-01-05", "1", "Laptop Pro", "2", "1299.99", "2599.98"],
        ["2023-01-10", "2", "Wireless Headphones", "5", "149.99", "749.95"],
        ["2023-01-15", "4", "Programming Guide", "10", "39.99", "399.90"],
        ["2023-01-20", "5", "Coffee Maker", "3", "89.99", "269.97"],
        ["2023-01-25", "6", "Bluetooth Speaker", "4", "79.99", "319.96"],
        ["2023-02-03", "2", "Wireless Headphones", "3", "149.99", "449.97"],
        ["2023-02-08", "3", "Smart Watch", "2", "249.99", "499.98"],
        ["2023-02-12", "4", "Programming Guide", "5", "39.99", "199.95"],
        ["2023-02-18", "1", "Laptop Pro", "1", "1299.99", "1299.99"],
        ["2023-02-25", "6", "Bluetooth Speaker", "2", "79.99", "159.98"]
    ]

    with open(sales_file_path, "w", newline="") as f:
        writer = csv.writer(f)
        for row in sales_data:
            writer.writerow(row)

    print(f"  Created {sales_filename}")

    # Customer data
    customer_filename = "customer_data.csv"
    customer_file_path = os.path.join(output_dir, customer_filename)

    customer_data = [
        ["customer_id", "name", "email", "signup_date", "loyalty_tier", "purchase_count"],
        ["C001", "John Doe", "john.doe@example.com", "2022-01-15", "Gold", "24"],
        ["C002", "Jane Smith", "jane.smith@example.com", "2022-02-20", "Silver", "16"],
        ["C003", "Robert Johnson", "robert.j@example.com", "2022-03-05", "Bronze", "8"],
        ["C004", "Emily Davis", "emily.d@example.com", "2022-04-10", "Gold", "31"],
        ["C005", "Michael Wilson", "michael.w@example.com", "2022-05-22", "Silver", "14"],
        ["C006", "Sarah Thompson", "sarah.t@example.com", "2022-06-18", "Bronze", "6"],
        ["C007", "David Miller", "david.m@example.com", "2022-07-30", "Gold", "28"],
        ["C008", "Jennifer White", "jennifer.w@example.com", "2022-08-12", "Silver", "12"]
    ]

    with open(customer_file_path, "w", newline="") as f:
        writer = csv.writer(f)
        for row in customer_data:
            writer.writerow(row)

    print(f"  Created {customer_filename}")

    return True

def main():
    parser = argparse.ArgumentParser(description="Create sample data files for AI Assistant testing")
    parser.add_argument("--output-dir", "-o", type=str, default="./data/host_files",
                        help="Output directory for sample files")
    parser.add_argument("--text", type=int, default=3,
                        help="Number of text files to create")
    parser.add_argument("--pdf", type=int, default=2,
                        help="Number of PDF files to create")
    parser.add_argument("--docx", type=int, default=2,
                        help="Number of DOCX files to create")
    parser.add_argument("--json", action="store_true",
                        help="Create sample JSON files")
    parser.add_argument("--csv", action="store_true",
                        help="Create sample CSV files")
    parser.add_argument("--all", action="store_true",
                        help="Create all sample file types")

    args = parser.parse_args()

    output_dir = args.output_dir

    # Create the output directory if it doesn't exist
    os.makedirs(output_dir, exist_ok=True)

    # Create sample files based on arguments
    if args.all or args.text > 0:
        create_sample_text_files(output_dir, args.text)

    if args.all or args.pdf > 0:
        create_sample_pdf_files(output_dir, args.pdf)

    if args.all or args.docx > 0:
        create_sample_docx_files(output_dir, args.docx)

    if args.all or args.json:
        create_sample_json_files(output_dir)

    if args.all or args.csv:
        create_sample_csv_files(output_dir)

    print(f"\nSample data creation complete! Files are available in {output_dir}")

    return 0

if __name__ == "__main__":
    sys.exit(main())